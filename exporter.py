"""
exporter.py — Export a pipeline result dict to a formatted Word document.

Produces a .docx with:
  - Cover page with ML prediction summary table
  - Full procurement briefing report (parsed from markdown headings)
  - Plausibility critique and analyst commentary
  - Raw numbers appendix (ML regression, bucket classifier)
  - Similar historical contracts appendix

Usage:
    from exporter import export_to_word
    export_to_word(result, "report.docx")
"""

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

from utils import fmt_dollar

# ── Style constants ────────────────────────────────────────────────────────────
HEADING_COLOR   = (31, 73, 125)     # dark blue
APPENDIX_COLOR  = (89, 89, 89)      # grey
HEADER_BG_HEX   = "1F497D"          # table header fill
ALT_ROW_BG_HEX  = "EEF2F7"         # alternating row fill
KEY_COL_BG_HEX  = "E8ECF3"         # kv-table label column fill
TABLE_STYLE     = "Table Grid"


# ── Helpers ───────────────────────────────────────────────────────────────────

def _set_heading_color(paragraph, rgb: tuple[int, int, int]):
    """Apply an RGB color to all runs in a paragraph."""
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*rgb)


def _add_horizontal_rule(doc: Document):
    """Insert a thin horizontal line (paragraph border)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)



def _parse_markdown_sections(text: str) -> list[tuple[str, str]]:
    """
    Split a markdown string into (heading, body) pairs.
    Handles ## and ### headings. Returns [(heading, body), ...].
    """
    sections = []
    current_heading = None
    current_body: list[str] = []

    for line in text.splitlines():
        m = re.match(r"^#{1,3}\s+(.+)", line)
        if m:
            if current_heading is not None:
                sections.append((current_heading, "\n".join(current_body).strip()))
            current_heading = m.group(1).strip()
            current_body = []
        else:
            if current_heading is not None:
                current_body.append(line)

    if current_heading is not None:
        sections.append((current_heading, "\n".join(current_body).strip()))

    return sections


def _set_cell_bg(cell, hex_color: str):
    """Set a table cell's background fill colour."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_padding(cell, top=60, bottom=60, left=108, right=108):
    """Set cell padding in twentieths of a point."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar  = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _parse_md_table(lines: list[str]) -> list[list[str]]:
    """Parse markdown table lines into a list of cell-lists, skipping separator rows."""
    rows = []
    for line in lines:
        stripped = line.strip()
        if re.match(r"^\|[-:\s|]+\|$", stripped):   # separator row
            continue
        cells = [c.strip() for c in stripped.strip("|").split("|")]
        rows.append(cells)
    return rows


def _write_md_table(doc: Document, lines: list[str]):
    """Render a block of markdown table lines as a styled Word table."""
    rows = _parse_md_table(lines)
    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    table  = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = TABLE_STYLE

    HEADER_FG  = RGBColor(255, 255, 255)

    for i, row in enumerate(rows):
        is_header = (i == 0)
        for j in range(n_cols):
            cell      = table.cell(i, j)
            cell_text = row[j] if j < len(row) else ""
            _set_cell_padding(cell)

            p = cell.paragraphs[0]
            p.clear()
            _add_inline_formatting(p, cell_text)

            if is_header:
                _set_cell_bg(cell, HEADER_BG_HEX)
                for run in p.runs:
                    run.bold            = True
                    run.font.color.rgb  = HEADER_FG
                    run.font.size       = Pt(9)
            else:
                if i % 2 == 0:
                    _set_cell_bg(cell, ALT_ROW_BG_HEX)
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacer after table


def _write_markdown_body(doc: Document, body: str):
    """
    Write a markdown body string into the document.
    Handles markdown tables, bullet points, numbered lists, bold, and plain text.
    """
    lines   = body.splitlines()
    i       = 0
    n       = len(lines)

    while i < n:
        line    = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # ── Markdown table block ───────────────────────────────────────────────
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i])
                i += 1
            _write_md_table(doc, table_lines)
            continue

        # ── Bullet point ───────────────────────────────────────────────────────
        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, stripped[2:])

        # ── Numbered list ──────────────────────────────────────────────────────
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, re.sub(r"^\d+\.\s+", "", stripped))

        # ── Plain paragraph ────────────────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, stripped)
            p.paragraph_format.space_after = Pt(4)

        i += 1


def _add_inline_formatting(paragraph, text: str):
    """Parse **bold** markers and add runs with correct formatting."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run      = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _add_kv_table(doc: Document, rows: list[tuple[str, str]]):
    """Two-column key/value table with a shaded key column."""
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = TABLE_STYLE

    for i, (key, val) in enumerate(rows):
        key_cell = table.cell(i, 0)
        val_cell = table.cell(i, 1)

        _set_cell_padding(key_cell)
        _set_cell_padding(val_cell)
        _set_cell_bg(key_cell, KEY_COL_BG_HEX)

        kp = key_cell.paragraphs[0]
        kp.clear()
        run       = kp.add_run(key)
        run.bold  = True
        run.font.size = Pt(9)

        vp = val_cell.paragraphs[0]
        vp.clear()
        run = vp.add_run(str(val))
        run.font.size = Pt(9)

    doc.add_paragraph()  # spacer


# ── Main export function ───────────────────────────────────────────────────────

def export_to_word(result: dict, output_path: str) -> str:
    """
    Write a pipeline result dict to a Word document.

    Args:
        result:      The dict returned by langchain_agents.graph.predict()
        output_path: Destination .docx path

    Returns:
        Absolute path to the written file.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Cover / header ────────────────────────────────────────────────────────
    title = doc.add_heading("Tender Price Prediction", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_heading_color(title, (31, 73, 125))  # dark blue

    sub = doc.add_paragraph("Procurement Briefing Report")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    date_p = doc.add_paragraph(datetime.now().strftime("%d %B %Y"))
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_p.runs[0].font.size = Pt(10)
    date_p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # ── Combined estimate summary box ─────────────────────────────────────────
    reg      = result.get("regression_prediction", {})
    bkt      = result.get("bucket_prediction", {})
    contract = result.get("contract", {})

    h = doc.add_heading("ML Price Prediction", level=2)
    _set_heading_color(h, HEADING_COLOR)

    summary_rows = [
        ("Point estimate",   fmt_dollar(reg.get("point_estimate_aud"))),
        ("90% CI lower",     fmt_dollar(reg.get("ci_low_90_aud"))),
        ("90% CI upper",     fmt_dollar(reg.get("ci_high_90_aud"))),
        ("Predicted bucket", bkt.get("predicted_bucket", "N/A")),
        ("Predicted range",  bkt.get("predicted_subrange", "N/A")),
        ("Bucket confidence", f"{bkt.get('bucket_probability', 'N/A')}"),
    ]
    _add_kv_table(doc, summary_rows)

    # ── Full report (parsed from markdown) ────────────────────────────────────
    report_text = result.get("report", "")
    if report_text:
        _add_horizontal_rule(doc)
        doc.add_paragraph()
        sections = _parse_markdown_sections(report_text)

        if sections:
            for heading, body in sections:
                h = doc.add_heading(heading, level=2)
                _set_heading_color(h, HEADING_COLOR)
                if body:
                    _write_markdown_body(doc, body)
                doc.add_paragraph()
        else:
            # Fallback: dump as plain text if no headings found
            doc.add_paragraph(report_text)

    # ── Plausibility critique ─────────────────────────────────────────────────
    critique = result.get("ml_critique", "")
    if critique:
        _add_horizontal_rule(doc)
        doc.add_paragraph()
        h = doc.add_heading("Plausibility Critique", level=1)
        _set_heading_color(h, HEADING_COLOR)
        _write_markdown_body(doc, critique)
        doc.add_paragraph()

    # ── Analyst commentary ────────────────────────────────────────────────────
    analysis = result.get("analysis", "")
    if analysis:
        _add_horizontal_rule(doc)
        doc.add_paragraph()
        h = doc.add_heading("Analyst Commentary", level=1)
        _set_heading_color(h, HEADING_COLOR)
        _write_markdown_body(doc, analysis)
        doc.add_paragraph()

    # ── Appendix: raw numbers ─────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    doc.add_paragraph()
    h = doc.add_heading("Appendix — Raw Model Outputs", level=1)
    _set_heading_color(h, APPENDIX_COLOR)

    doc.add_heading("Contract Features", level=3)
    if contract:
        _add_kv_table(doc, [(k, str(v)) for k, v in contract.items()])

    doc.add_heading("Statistical Model — Regression", level=3)
    if reg:
        reg_rows = [
            ("Point estimate",       fmt_dollar(reg.get("point_estimate_aud"))),
            ("90% CI lower",         fmt_dollar(reg.get("ci_low_90_aud"))),
            ("90% CI upper",         fmt_dollar(reg.get("ci_high_90_aud"))),
        ]
        _add_kv_table(doc, reg_rows)

    doc.add_heading("Statistical Model — Bucket Classifier", level=3)
    if bkt:
        bkt_rows = [
            ("Predicted bucket",     bkt.get("predicted_bucket", "N/A")),
            ("Bucket probability",   f"{bkt.get('bucket_probability', 'N/A')}"),
            ("Predicted sub-range",  bkt.get("predicted_subrange", "N/A")),
            ("Sub-range probability",f"{bkt.get('subrange_probability', 'N/A')}"),
            ("Sub-range low",        fmt_dollar(bkt.get("subrange_low_aud"))),
            ("Sub-range high",       fmt_dollar(bkt.get("subrange_high_aud"))),
        ]
        _add_kv_table(doc, bkt_rows)

    # ── Similar contracts ─────────────────────────────────────────────────────
    similar = result.get("similar_contracts", [])
    if similar:
        _add_horizontal_rule(doc)
        doc.add_paragraph()
        h = doc.add_heading("Appendix — Similar Historical Contracts", level=1)
        _set_heading_color(h, APPENDIX_COLOR)

        for i, c in enumerate(similar, 1):
            doc.add_heading(f"Contract {i}", level=3)
            rows = [(k, str(v)) for k, v in c.items()]
            _add_kv_table(doc, rows)

    # ── Model performance comparison (if training has been run) ───────────────
    _add_model_comparison_appendix(doc, reg.get("model_key", "xgboost"))

    # ── Save ──────────────────────────────────────────────────────────────────
    out = Path(output_path).resolve()
    doc.save(str(out))
    return str(out)


def _add_model_comparison_appendix(doc: Document, active_model_key: str) -> None:
    """Add a Model Performance Comparison appendix if comparison data exists."""
    try:
        from ml_evaluation.evaluator import MultiModelEvaluator
        comparison = MultiModelEvaluator.load_comparison()
    except Exception:
        return

    rows_ok = [r for r in comparison if r.get("status") == "ok"]
    if not rows_ok:
        return

    rows_ok.sort(key=lambda r: r.get("r2", 0), reverse=True)

    _add_horizontal_rule(doc)
    doc.add_paragraph()
    h = doc.add_heading("Appendix — ML Model Performance Comparison", level=1)
    _set_heading_color(h, APPENDIX_COLOR)

    note = doc.add_paragraph(
        f"Prediction generated using: {rows_ok[0]['display_name'] if not any(r['model_key'] == active_model_key for r in rows_ok) else next(r['display_name'] for r in rows_ok if r['model_key'] == active_model_key)}.  "
        "All models trained on the same 80/20 train/test split.  "
        "Metrics are in log₁p space unless noted."
    )
    note.runs[0].font.size = Pt(9)
    note.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    doc.add_paragraph()

    # Build table: header + one row per model
    headers = ["Model", "R²", "RMSE (log)", "MAE (log)", "MAE ($)", "≤50% acc.", "Time (s)"]
    table = doc.add_table(rows=1 + len(rows_ok), cols=len(headers))
    table.style = TABLE_STYLE

    # Header row
    hdr_row = table.rows[0]
    for j, hdr in enumerate(headers):
        cell = hdr_row.cells[j]
        _set_cell_padding(cell)
        _set_cell_bg(cell, HEADER_BG_HEX)
        p = cell.paragraphs[0]
        p.clear()
        run = p.add_run(hdr)
        run.bold = True
        run.font.size  = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)

    # Data rows
    for i, row in enumerate(rows_ok):
        is_active = row["model_key"] == active_model_key
        tr = table.rows[i + 1]
        if i % 2 == 0:
            bg = ALT_ROW_BG_HEX
        else:
            bg = "FFFFFF"

        values = [
            row["display_name"] + (" ★" if is_active else ""),
            f"{row['r2']:.4f}",
            f"{row['rmse_log']:.4f}",
            f"{row['mae_log']:.4f}",
            f"${row.get('mae_dollar', 0):>12,.0f}" if row.get("mae_dollar") else "—",
            f"{row.get('within_50pct', 0):.1f}%",
            f"{row.get('train_time_s', 0):.1f}s",
        ]
        for j, val in enumerate(values):
            cell = tr.cells[j]
            _set_cell_padding(cell)
            _set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(val)
            run.font.size = Pt(9)
            if is_active:
                run.bold = True

    doc.add_paragraph()
