"""
evaluation/run_evaluation.py — Research evaluation for the Tender Price Prediction system.

Generates a Word document with three evaluation sections:
  1. ML Model Evaluation   — regression metrics on held-out test set
  2. RAG Evaluation        — UNSPSC retrieval accuracy
  3. End-to-End Evaluation — full pipeline: NL description → predicted price vs actual

Requirements:
  - tenders_export.xlsx must exist in project root
  - Trained models must exist in models/
  - Domain RAG index must be built (domain_rag_store/)
  - App must be running at localhost:8000 for end-to-end evaluation

Usage:
    python evaluation/run_evaluation.py --data tenders_export.xlsx
    python evaluation/run_evaluation.py --data tenders_export.xlsx --e2e-samples 20 --rag-samples 100
"""

import argparse
import json
import os
import sys
import time

import numpy as np
import pandas as pd
import requests

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from config import MODELS_DIR, UNSPSC_FILE


# ── Helpers ────────────────────────────────────────────────────────────────────

def fmt_dollar(val: float) -> str:
    if val >= 1_000_000:
        return f"${val / 1_000_000:,.2f}M"
    elif val >= 1_000:
        return f"${val / 1_000:,.1f}K"
    return f"${val:,.0f}"


def load_unspsc_lookup() -> dict[str, str]:
    """Build a code → title lookup from the UNSPSC Excel file."""
    try:
        df = pd.read_excel(UNSPSC_FILE, dtype=str)
        col_map = {c.lower().strip(): c for c in df.columns}
        code_col  = col_map.get("code") or col_map.get("commodity code") or df.columns[0]
        title_col = col_map.get("title") or col_map.get("commodity title") or df.columns[1]
        df[code_col] = df[code_col].str.strip().str.split(".").str[0]
        return dict(zip(df[code_col], df[title_col]))
    except Exception as exc:
        print(f"[WARN] Could not load UNSPSC lookup: {exc}")
        return {}


def generate_nl_prompt(row: dict, unspsc_lookup: dict) -> str:
    """Convert ground-truth contract fields to a natural language procurement description."""
    method      = str(row.get("procurement_method", "open tender")).lower()
    disposition = str(row.get("disposition", "contract notice")).lower()
    publisher   = str(row.get("publisher_name", "a government agency"))
    duration    = row.get("duration_days")

    # Derive year/quarter from contract_start if pre-computed fields are absent
    year    = row.get("contract_start_year")
    quarter = row.get("contract_start_quarter")
    if not year:
        cs = pd.to_datetime(row.get("contract_start"), errors="coerce")
        year    = cs.year    if pd.notna(cs) else 2024
        quarter = cs.quarter if pd.notna(cs) else 2

    # UNSPSC title
    code = str(row.get("category_code", "")).split(".")[0].split(".")[0]
    category_title = unspsc_lookup.get(code, "goods and services").lower()

    # Duration string
    if duration and not pd.isna(duration):
        months = round(float(duration) / 30)
        duration_str = f"approximately {months} months" if months < 24 else f"approximately {round(months/12, 1)} years"
    else:
        duration_str = "an unspecified duration"

    # Start period
    quarter_map = {1: "early", 2: "mid", 3: "mid-late", 4: "late"}
    period = f"{quarter_map.get(int(quarter) if quarter and not pd.isna(quarter) else 2, 'mid')}-{int(year)}"

    return (
        f"I need to procure {category_title} for {publisher} via {method}. "
        f"It will be a {disposition}, {duration_str} duration, starting {period}."
    )


# ── 1. ML Evaluation ───────────────────────────────────────────────────────────

def run_ml_evaluation(data_path: str) -> dict:
    """Evaluate ML models on held-out test set."""
    print("\n[1/3] Running ML evaluation …")

    from ml_evaluation.evaluator import MultiModelEvaluator, get_active_model, ordinal_encode_split
    from pipeline.data_processor import DataProcessor
    from pipeline.regressor import Regressor

    # Load stored comparison metrics
    comparison = MultiModelEvaluator.load_comparison()
    active_key = get_active_model()

    # Load and split data
    dp = DataProcessor(verbose=False)
    dp._load_ohe_schema()

    df_raw = pd.read_excel(data_path)
    df_raw = df_raw[df_raw["value"] > 0].dropna(subset=["value"])
    df_raw["log_value"] = np.log1p(df_raw["value"])

    # Rebuild test set with same random_state=42
    from sklearn.model_selection import train_test_split
    context = {"data_path": data_path}
    dp.run(context)
    X_train = context["X_train"]
    X_test  = context["X_test"]
    y_train = context["y_train"]
    y_test  = context["y_test"]
    raw_df  = context["raw_df"]

    # Predictions from active model
    regressor = Regressor(model_key=active_key, verbose=False)
    regressor._load()

    X_tr_enc, X_te_enc, encoder, cat_cols = ordinal_encode_split(X_train, X_test)

    from ml_evaluation.model_registry import MODEL_REGISTRY
    spec = MODEL_REGISTRY[active_key]
    X_eval = X_test if spec["native_categorical"] else X_te_enc

    y_pred_log = regressor.model.predict(X_eval)
    y_true_d   = np.expm1(y_test.values)
    y_pred_d   = np.expm1(y_pred_log)

    # Within-range accuracy
    ratios = np.abs(y_pred_d - y_true_d) / (y_true_d + 1)
    within_25 = float(np.mean(ratios < 0.25) * 100)
    within_50 = float(np.mean(ratios < 0.50) * 100)
    within_2x = float(np.mean((y_pred_d / (y_true_d + 1) <= 2) & (y_true_d / (y_pred_d + 1) <= 2)) * 100)

    # Median baseline (predict median value for every contract)
    median_val  = np.median(y_true_d)
    base_ratios = np.abs(median_val - y_true_d) / (y_true_d + 1)
    baseline_within_50 = float(np.mean(base_ratios < 0.50) * 100)

    # Category median baseline
    test_idx = X_test.index  # X preserves original df index; y uses reset .values index
    if "category_code" in raw_df.columns:
        cat_medians = raw_df.groupby("category_code")["value"].median()
        cat_preds   = raw_df.loc[test_idx, "category_code"].map(cat_medians).fillna(median_val)
        cat_ratios  = np.abs(cat_preds.values - y_true_d) / (y_true_d + 1)
        cat_baseline_within_50 = float(np.mean(cat_ratios < 0.50) * 100)
    else:
        cat_baseline_within_50 = baseline_within_50

    active_metrics = next((r for r in comparison if r.get("model_key") == active_key and r.get("status") == "ok"), {})

    # ── Price bucket breakdown ────────────────────────────────────────────────
    price_buckets = {
        "Small (<$100K)":       (y_true_d < 100_000),
        "Medium ($100K–$1M)":   (y_true_d >= 100_000) & (y_true_d < 1_000_000),
        "Large (>$1M)":         (y_true_d >= 1_000_000),
    }
    price_breakdown = {}
    for label, mask in price_buckets.items():
        if mask.sum() == 0:
            price_breakdown[label] = {"n": 0, "within_50": 0.0, "within_2x": 0.0}
            continue
        r = np.abs(y_pred_d[mask] - y_true_d[mask]) / (y_true_d[mask] + 1)
        price_breakdown[label] = {
            "n":         int(mask.sum()),
            "within_50": round(float(np.mean(r < 0.50) * 100), 1),
            "within_2x": round(float(np.mean(
                (y_pred_d[mask] / (y_true_d[mask] + 1) <= 2) &
                (y_true_d[mask] / (y_pred_d[mask] + 1) <= 2)
            ) * 100), 1),
        }

    # ── Time bucket breakdown ─────────────────────────────────────────────────
    time_breakdown = {}
    if "contract_start" in raw_df.columns:
        raw_df["_year"] = pd.to_datetime(raw_df["contract_start"], errors="coerce").dt.year
        test_years = raw_df.loc[test_idx, "_year"].values
        time_buckets = {
            "2020–2022": (test_years >= 2020) & (test_years <= 2022),
            "2022–2024": (test_years >= 2022) & (test_years <= 2024),
            "2024–2026": (test_years >= 2024) & (test_years <= 2026),
        }
        for label, mask in time_buckets.items():
            if mask.sum() == 0:
                time_breakdown[label] = {"n": 0, "within_50": 0.0, "within_2x": 0.0}
                continue
            r = np.abs(y_pred_d[mask] - y_true_d[mask]) / (y_true_d[mask] + 1)
            time_breakdown[label] = {
                "n":         int(mask.sum()),
                "within_50": round(float(np.mean(r < 0.50) * 100), 1),
                "within_2x": round(float(np.mean(
                    (y_pred_d[mask] / (y_true_d[mask] + 1) <= 2) &
                    (y_true_d[mask] / (y_pred_d[mask] + 1) <= 2)
                ) * 100), 1),
            }

    # ── KNN range coverage ────────────────────────────────────────────────────
    print("  Evaluating KNN price range coverage …")
    from rag.knn_retriever import search_contracts as search_similar_contracts
    knn_in_range = 0
    knn_samples  = min(500, len(X_test))
    knn_sample_idx = np.random.RandomState(42).choice(len(X_test), knn_samples, replace=False)

    for si in knn_sample_idx:
        actual = float(y_true_d[si])
        row    = raw_df.iloc[list(raw_df.index).index(X_test.index[si])] if X_test.index[si] in raw_df.index else None
        if row is None:
            continue
        contract_dict = {
            "procurement_method":  str(row.get("procurement_method", "")),
            "disposition":         str(row.get("disposition", "")),
            "publisher_gov_type":  str(row.get("publisher_gov_type", "fed")),
            "category_code":       str(row.get("category_code", "")).split(".")[0],
            "parent_category_code": str(row.get("parent_category_code", "")).split(".")[0],
            "publisher_cofog_level": str(row.get("publisher_cofog_level", "")),
            "publisher_name":      str(row.get("publisher_name", "")),
            "duration_days":       float(row.get("duration_days", 330)) if pd.notna(row.get("duration_days")) else 330.0,
        }
        try:
            neighbors = search_similar_contracts(contract_dict, n_results=5)
            if neighbors:
                vals     = [float(n["value"]) for n in neighbors if n.get("value")]
                knn_min  = min(vals)
                knn_max  = max(vals)
                if knn_min <= actual <= knn_max:
                    knn_in_range += 1
        except Exception:
            pass

    knn_coverage = round(knn_in_range / knn_samples * 100, 1)
    print(f"  KNN range coverage (actual within KNN min–max): {knn_coverage}%  (n={knn_samples})")

    print(f"  Active model: {active_key}  R²={active_metrics.get('r2', 'N/A')}  Within-50%={within_50:.1f}%")
    for label, stats in price_breakdown.items():
        print(f"    {label}: n={stats['n']}  within-50%={stats['within_50']}%")

    return {
        "active_key":             active_key,
        "active_metrics":         active_metrics,
        "all_models":             [r for r in comparison if r.get("status") == "ok"],
        "within_25pct":           within_25,
        "within_50pct":           within_50,
        "within_2x":              within_2x,
        "baseline_within_50":     baseline_within_50,
        "cat_baseline_within_50": cat_baseline_within_50,
        "n_test":                 len(y_test),
        "price_breakdown":        price_breakdown,
        "time_breakdown":         time_breakdown,
        "knn_coverage":           knn_coverage,
        "knn_samples":            knn_samples,
    }


# ── 2. RAG Evaluation ──────────────────────────────────────────────────────────

def run_rag_evaluation(data_path: str, n_samples: int = 100) -> dict:
    """Evaluate UNSPSC retrieval accuracy using domain RAG."""
    print(f"\n[2/3] Running RAG evaluation (n={n_samples}) …")

    from tools.domain_tools import lookup_procurement_codes

    unspsc_lookup = load_unspsc_lookup()
    if not unspsc_lookup:
        return {"error": "UNSPSC lookup not available"}

    df = pd.read_excel(data_path)
    df = df[df["value"] > 0].dropna(subset=["value", "category_code"])
    df["category_code_str"] = df["category_code"].astype(str).str.split(".").str[0]
    df = df[df["category_code_str"].isin(unspsc_lookup)]

    sample = df.sample(min(n_samples, len(df)), random_state=42)

    top1_exact = 0
    top3_exact = 0
    top1_segment = 0
    results = []

    for i, (_, row) in enumerate(sample.iterrows()):
        actual_code = row["category_code_str"]
        title = unspsc_lookup.get(actual_code, "")
        if not title or title.lower() in ("nan", ""):
            continue

        try:
            raw = lookup_procurement_codes.invoke({"description": title, "field": "category_code"})
            matches = json.loads(raw)
            if isinstance(matches, list) and matches:
                top1_code = str(matches[0].get("category_code", "")).split(".")[0]
                top3_codes = [str(m.get("category_code", "")).split(".")[0] for m in matches[:3]]

                exact_top1 = top1_code == actual_code
                exact_top3 = actual_code in top3_codes
                seg_top1   = top1_code[:2] == actual_code[:2]

                if exact_top1: top1_exact += 1
                if exact_top3: top3_exact += 1
                if seg_top1:   top1_segment += 1

                results.append({
                    "actual_code":    actual_code,
                    "actual_title":   title,
                    "top1_code":      top1_code,
                    "top1_title":     matches[0].get("name", ""),
                    "exact_top1":     exact_top1,
                    "exact_top3":     exact_top3,
                    "segment_match":  seg_top1,
                    "similarity":     matches[0].get("similarity", 0),
                })

            if (i + 1) % 10 == 0:
                print(f"  {i+1}/{len(sample)} — top-1: {top1_exact/(i+1)*100:.1f}%")

        except Exception as exc:
            print(f"  [WARN] RAG error for code {actual_code}: {exc}")

    n = len(results)
    return {
        "n_samples":        n,
        "top1_exact_pct":   round(top1_exact / n * 100, 1) if n else 0,
        "top3_exact_pct":   round(top3_exact / n * 100, 1) if n else 0,
        "top1_segment_pct": round(top1_segment / n * 100, 1) if n else 0,
        "results":          results[:20],  # store first 20 for report table
    }


# ── 3. End-to-End Evaluation — field extraction accuracy ──────────────────────

def _field_match(gt_val, extracted_val, field: str) -> str:
    """Return 'exact', 'partial', or 'miss' for a single field comparison."""
    if extracted_val is None or str(extracted_val).lower() in ("unknown", "none", ""):
        return "miss"
    gt  = str(gt_val).strip().lower().split(".")[0]
    ext = str(extracted_val).strip().lower().split(".")[0]
    if field == "duration_days":
        try:
            return "exact" if abs(float(ext) - float(gt)) / (float(gt) + 1) < 0.20 else "miss"
        except Exception:
            return "miss"
    if field == "category_code":
        if ext == gt:               return "exact"
        if ext[:2] == gt[:2]:       return "partial"   # same UNSPSC segment
        return "miss"
    return "exact" if ext == gt else "miss"


def run_e2e_evaluation(data_path: str, n_samples: int = 50, api_url: str = "http://localhost:8000") -> dict:
    """Evaluate field extraction accuracy: NL description → agent → extracted fields vs ground truth."""
    print(f"\n[3/3] Running E2E field extraction evaluation (n={n_samples}) …")

    unspsc_lookup = load_unspsc_lookup()

    df = pd.read_excel(data_path)
    df = df[df["value"] > 0].dropna(subset=["value", "category_code", "publisher_name", "procurement_method"])
    df["contract_start"] = pd.to_datetime(df["contract_start"], errors="coerce")
    df["contract_end"]   = pd.to_datetime(df["contract_end"],   errors="coerce")
    df["duration_days"]  = (df["contract_end"] - df["contract_start"]).dt.days
    df = df[df["duration_days"].notna() & (df["duration_days"] > 0)]

    print(f"  E2E pool (complete features): {len(df):,} contracts")
    sample = df.sample(min(n_samples, len(df)), random_state=7)

    FIELDS = ["procurement_method", "disposition", "publisher_gov_type",
              "category_code", "publisher_cofog_level", "publisher_name", "duration_days"]

    field_counts  = {f: {"exact": 0, "partial": 0, "miss": 0} for f in FIELDS}
    results = []

    for i, (_, row) in enumerate(sample.iterrows()):
        prompt = generate_nl_prompt(row.to_dict(), unspsc_lookup)
        extracted = None

        try:
            resp = requests.post(f"{api_url}/api/chat", json={"message": prompt}, timeout=180)
            resp.raise_for_status()
            data = resp.json()
            session_id = data.get("session_id")
            extracted  = data.get("extracted_contract")

            if not extracted:
                resp2 = requests.post(
                    f"{api_url}/api/chat",
                    json={"message": "please run the prediction now", "session_id": session_id},
                    timeout=180,
                )
                resp2.raise_for_status()
                data2    = resp2.json()
                extracted = data2.get("extracted_contract")

            if session_id:
                requests.delete(f"{api_url}/api/session/{session_id}", timeout=10)

            if extracted:
                row_result = {"prompt": prompt[:120], "fields": {}}
                for f in FIELDS:
                    gt_val  = row.get(f)
                    ext_val = extracted.get(f)
                    match   = _field_match(gt_val, ext_val, f)
                    field_counts[f][match] += 1
                    row_result["fields"][f] = {
                        "ground_truth": str(gt_val)[:40],
                        "extracted":    str(ext_val)[:40],
                        "match":        match,
                    }
                results.append(row_result)
                n_correct = sum(1 for f in FIELDS if row_result["fields"][f]["match"] == "exact")
                print(f"  {i+1}/{len(sample)} — {n_correct}/{len(FIELDS)} fields correct  |  prompt: {prompt[:80]}…")
            else:
                results.append({"prompt": prompt[:120], "error": "No extraction returned"})
                print(f"  {i+1}/{len(sample)} — no extraction")

            time.sleep(2)

        except Exception as exc:
            print(f"  [WARN] E2E error for contract {i+1}: {exc}")
            results.append({"prompt": prompt[:80], "error": str(exc)})

    n = len([r for r in results if "fields" in r])

    # Per-field accuracy
    field_accuracy = {}
    for f in FIELDS:
        counts = field_counts[f]
        total  = counts["exact"] + counts["partial"] + counts["miss"]
        field_accuracy[f] = {
            "exact_pct":   round(counts["exact"]   / total * 100, 1) if total else 0,
            "partial_pct": round(counts["partial"] / total * 100, 1) if total else 0,
            "miss_pct":    round(counts["miss"]    / total * 100, 1) if total else 0,
        }

    overall_exact = round(
        sum(field_counts[f]["exact"] for f in FIELDS) / (n * len(FIELDS)) * 100, 1
    ) if n else 0

    print(f"\n  Overall field extraction accuracy: {overall_exact}%")
    for f, acc in field_accuracy.items():
        print(f"    {f}: exact={acc['exact_pct']}%  partial={acc['partial_pct']}%")

    return {
        "n_samples":       len(results),
        "n_successful":    n,
        "overall_exact":   overall_exact,
        "field_accuracy":  field_accuracy,
        "results":         results[:20],
    }


# ── Report Generation ──────────────────────────────────────────────────────────

def generate_report(ml: dict, rag: dict, e2e: dict, output_path: str) -> None:
    """Generate Word report. If ML+RAG both skipped and file exists, append pipeline section only."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    append_mode = ml.get("skipped") and rag.get("skipped") and os.path.exists(output_path)

    if append_mode:
        doc = Document(output_path)
        doc.add_page_break()
        print(f"  Appending to existing report: {output_path}")
    else:
        doc = Document()
        title = doc.add_heading("Tender Price Prediction — Evaluation Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(
            "This report evaluates the Tender Price Prediction system across three dimensions: "
            "ML model performance, domain RAG retrieval accuracy, and full pipeline accuracy."
        )
        doc.add_paragraph()

        # ── Section 1: ML Evaluation ──────────────────────────────────────────
        doc.add_heading("1. ML Model Evaluation", level=1)
        doc.add_paragraph(
            "The ML pipeline was evaluated on a held-out test set (20% of the full dataset, "
            "random_state=42). The target variable is the contract awarded value in AUD."
        )

        doc.add_heading("1.1 Model Comparison", level=2)
        all_models = ml.get("all_models", [])
        if all_models:
            headers = ["Model", "R²", "RMSE (log)", "MAE (log)", "MAE ($)", "Within 50%", "Train time (s)"]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for m in all_models:
                rc = table.add_row().cells
                rc[0].text = m.get("display_name", m.get("model_key", ""))
                rc[1].text = str(m.get("r2", ""))
                rc[2].text = str(m.get("rmse_log", ""))
                rc[3].text = str(m.get("mae_log", ""))
                rc[4].text = f"${m.get('mae_dollar', 0):,.0f}" if m.get("mae_dollar") else ""
                rc[5].text = f"{m.get('within_50pct', '')}%"
                rc[6].text = str(m.get("train_time_s", ""))

        doc.add_paragraph()
        doc.add_heading("1.2 Within-Range Accuracy (Active Model)", level=2)
        n_test = ml.get("n_test")
        n_test_str = f"{n_test:,}" if isinstance(n_test, int) else "N/A"
        doc.add_paragraph(
            f"Active model: {ml.get('active_key', 'N/A')}. Test set size: {n_test_str} contracts."
        )

        acc_table = doc.add_table(rows=1, cols=4)
        acc_table.style = "Table Grid"
        for i, h in enumerate(["Metric", "Active Model", "Global Median Baseline", "Category Median Baseline"]):
            acc_table.rows[0].cells[i].text = h
            acc_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for rd in [
            ("Within 25% of actual", f"{ml.get('within_25pct', 0):.1f}%", "—", "—"),
            ("Within 50% of actual", f"{ml.get('within_50pct', 0):.1f}%", f"{ml.get('baseline_within_50', 0):.1f}%", f"{ml.get('cat_baseline_within_50', 0):.1f}%"),
            ("Within 2× of actual",  f"{ml.get('within_2x', 0):.1f}%",    "—", "—"),
        ]:
            rc = acc_table.add_row().cells
            for i, v in enumerate(rd):
                rc[i].text = v

        doc.add_paragraph()

        # Price bucket breakdown
        price_breakdown = ml.get("price_breakdown", {})
        if price_breakdown:
            doc.add_heading("1.3 Accuracy by Contract Size", level=2)
            pt = doc.add_table(rows=1, cols=4)
            pt.style = "Table Grid"
            for i, h in enumerate(["Price Bucket", "n", "Within 50%", "Within 2×"]):
                pt.rows[0].cells[i].text = h
                pt.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for label, stats in price_breakdown.items():
                rc = pt.add_row().cells
                rc[0].text = label
                rc[1].text = str(stats["n"])
                rc[2].text = f"{stats['within_50']}%"
                rc[3].text = f"{stats['within_2x']}%"
            doc.add_paragraph()

        # Time bucket breakdown
        time_breakdown = ml.get("time_breakdown", {})
        if time_breakdown:
            doc.add_heading("1.4 Accuracy by Contract Year", level=2)
            tt = doc.add_table(rows=1, cols=4)
            tt.style = "Table Grid"
            for i, h in enumerate(["Period", "n", "Within 50%", "Within 2×"]):
                tt.rows[0].cells[i].text = h
                tt.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for label, stats in time_breakdown.items():
                rc = tt.add_row().cells
                rc[0].text = label
                rc[1].text = str(stats["n"])
                rc[2].text = f"{stats['within_50']}%"
                rc[3].text = f"{stats['within_2x']}%"
        # KNN range coverage
        doc.add_heading("1.5 KNN Price Range Coverage", level=2)
        doc.add_paragraph(
            "The KNN retriever finds the 5 most similar historical contracts in the training set "
            "using the same feature space as the ML model. It provides a price range (min–max) "
            "rather than a point estimate, offering a plausibility check on the regression output."
        )
        knn_table = doc.add_table(rows=1, cols=2)
        knn_table.style = "Table Grid"
        for i, h in enumerate(["Metric", "Value"]):
            knn_table.rows[0].cells[i].text = h
            knn_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for kr in [
            ("Contracts evaluated",                   str(ml.get("knn_samples", 0))),
            ("Actual value within KNN min–max range", f"{ml.get('knn_coverage', 0)}%"),
            ("Neighbours retrieved per contract",     "5"),
        ]:
            rc = knn_table.add_row().cells
            rc[0].text = kr[0]; rc[1].text = kr[1]
        doc.add_paragraph()

        # ── Section 2: RAG Evaluation ─────────────────────────────────────────
        doc.add_heading("2. Domain RAG Evaluation", level=1)
        doc.add_paragraph(
            "The Domain RAG index resolves natural language descriptions into structured codes "
            "across three domains:\n"
            "  • UNSPSC commodity codes — 8-digit codes identifying what is being procured "
            "(e.g. 'legal services' → 80120000)\n"
            "  • COFOG functional classification — maps agency function to government expenditure "
            "category (e.g. 'Department of Defence' → General public services)\n"
            "  • AusTender valid values — valid procurement methods, disposition types, and "
            "publisher government types used as structured inputs to the ML model.\n\n"
            "The RAG index was evaluated specifically on UNSPSC retrieval, as this is the most "
            "critical lookup (category_code is the strongest ML predictor)."
        )

        if "error" in rag:
            doc.add_paragraph(f"Evaluation skipped: {rag['error']}")
        else:
            doc.add_heading("2.1 Retrieval Accuracy", level=2)
            doc.add_paragraph(f"Sample size: {rag.get('n_samples', 0)} contracts.")

            rag_table = doc.add_table(rows=1, cols=2)
            rag_table.style = "Table Grid"
            for i, h in enumerate(["Metric", "Accuracy"]):
                rag_table.rows[0].cells[i].text = h
                rag_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for rr in [
                ("Top-1 exact match",                    f"{rag.get('top1_exact_pct', 0)}%"),
                ("Top-3 exact match",                    f"{rag.get('top3_exact_pct', 0)}%"),
                ("Top-1 segment match (first 2 digits)", f"{rag.get('top1_segment_pct', 0)}%"),
            ]:
                r = rag_table.add_row().cells
                r[0].text = rr[0]; r[1].text = rr[1]

            doc.add_paragraph()
            doc.add_heading("2.2 Sample Results (first 10)", level=2)
            sample_results = rag.get("results", [])[:10]
            if sample_results:
                t = doc.add_table(rows=1, cols=4)
                t.style = "Table Grid"
                for i, hh in enumerate(["Actual Code", "Actual Title", "Retrieved Code", "Match"]):
                    t.rows[0].cells[i].text = hh
                    t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                for sr in sample_results:
                    rc = t.add_row().cells
                    rc[0].text = sr.get("actual_code", "")
                    rc[1].text = str(sr.get("actual_title", ""))[:40]
                    rc[2].text = sr.get("top1_code", "")
                    rc[3].text = "✓" if sr.get("exact_top1") else ("~" if sr.get("segment_match") else "✗")

        doc.add_paragraph()

    # ── Section 3: End-to-End Evaluation ─────────────────────────────────────
    doc.add_heading("3. End-to-End Evaluation — Agent Field Extraction Accuracy", level=1)
    doc.add_paragraph(
        "The conversational agent is responsible for understanding a free-form natural language "
        "procurement description and mapping it to the 7 structured pre-award features required "
        "by the ML model. The agent handles two types of field extraction:\n\n"
        "  • Direct extraction (agent language understanding): procurement_method, disposition, "
        "publisher_name, publisher_gov_type, duration_days — these are inferred directly from "
        "the natural language text without external lookup.\n\n"
        "  • RAG-assisted extraction: category_code (UNSPSC), publisher_cofog_level (COFOG) — "
        "the agent calls the domain RAG lookup tool to resolve these into structured codes, "
        "as they require domain-specific knowledge the agent cannot reliably infer alone.\n\n"
        "This section evaluates how accurately the agent extracts each field, comparing the "
        "values it passes to predict_contract against the ground-truth values from the dataset. "
        "Category code is scored as 'partial' if the UNSPSC segment (first 2 digits) matches."
    )

    doc.add_heading("3.1 Overall Accuracy", level=2)
    doc.add_paragraph(
        f"Contracts evaluated: {e2e.get('n_samples', 0)}. "
        f"Successful extractions: {e2e.get('n_successful', 0)}. "
        f"Overall exact field match: {e2e.get('overall_exact', 0)}%."
    )

    doc.add_heading("3.2 Per-Field Accuracy", level=2)
    field_accuracy = e2e.get("field_accuracy", {})
    if field_accuracy:
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for i, hh in enumerate(["Field", "Exact Match", "Partial Match", "Miss"]):
            t.rows[0].cells[i].text = hh
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for field, acc in field_accuracy.items():
            rc = t.add_row().cells
            rc[0].text = field
            rc[1].text = f"{acc['exact_pct']}%"
            rc[2].text = f"{acc['partial_pct']}%"
            rc[3].text = f"{acc['miss_pct']}%"

    doc.add_paragraph()

    doc.add_heading("3.3 Sample Extractions (first 5)", level=2)
    sample_results = [r for r in e2e.get("results", []) if "fields" in r][:5]
    for sr in sample_results:
        doc.add_paragraph(f"Prompt: {sr.get('prompt', '')}", style="Intense Quote")
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        for i, hh in enumerate(["Field", "Ground Truth", "Extracted"]):
            t.rows[0].cells[i].text = hh
            t.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for field, fdata in sr["fields"].items():
            rc = t.add_row().cells
            rc[0].text = field
            rc[1].text = fdata["ground_truth"]
            rc[2].text = f"{fdata['extracted']} ({'✓' if fdata['match'] == 'exact' else '~' if fdata['match'] == 'partial' else '✗'})"
        doc.add_paragraph()

    doc.save(output_path)
    print(f"\n[✓] Report saved → {output_path}")


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Run evaluation and generate report")
    parser.add_argument("--data",         default="tenders_export.xlsx", help="Path to tenders Excel file")
    parser.add_argument("--output",       default="evaluation_report.docx", help="Output Word document path")
    parser.add_argument("--rag-samples",  type=int, default=100, help="Number of contracts for RAG evaluation")
    parser.add_argument("--e2e-samples",  type=int, default=50,  help="Number of contracts for E2E evaluation")
    parser.add_argument("--api-url",      default="http://localhost:8000", help="App URL for E2E evaluation")
    parser.add_argument("--skip-e2e",     action="store_true", help="Skip end-to-end evaluation")
    parser.add_argument("--skip-ml",      action="store_true", help="Skip ML evaluation")
    parser.add_argument("--skip-rag",     action="store_true", help="Skip RAG evaluation")
    args = parser.parse_args()

    if not os.path.exists(args.data):
        print(f"Error: data file not found: {args.data}")
        sys.exit(1)

    print(f"Starting evaluation — data: {args.data}")
    t_start = time.time()

    if args.skip_ml:
        ml_results = {"active_key": "N/A", "skipped": True}
        print("\n[1/3] ML evaluation skipped (--skip-ml)")
    else:
        ml_results = run_ml_evaluation(args.data)

    if args.skip_rag:
        rag_results = {"n_samples": 0, "n_successful": 0, "results": [], "skipped": True}
        print("\n[2/3] RAG evaluation skipped (--skip-rag)")
    else:
        rag_results = run_rag_evaluation(args.data, n_samples=args.rag_samples)

    if args.skip_e2e:
        e2e_results = {"n_samples": 0, "n_successful": 0, "results": [], "skipped": True}
        print("\n[3/3] End-to-end evaluation skipped (--skip-e2e)")
    else:
        e2e_results = run_e2e_evaluation(args.data, n_samples=args.e2e_samples, api_url=args.api_url)

    generate_report(ml_results, rag_results, e2e_results, args.output)
    print(f"\nTotal time: {(time.time() - t_start) / 60:.1f} minutes")


if __name__ == "__main__":
    main()
