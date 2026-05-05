"""
generate_eda_report.py — Build a Word document for the Tendertrace EDA.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime, os

OUT = "eda_report.docx"
IMG = "eda_output"

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color="1E3A5F"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(color)
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in p.runs:
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(40, 40, 40)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(30, 58, 95)
        r2 = p.add_run(text)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = RGBColor(40, 40, 40)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(40, 40, 40)
    return p

def add_figure(doc, img_path, caption, width=6.0):
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()  # spacer

def add_stat_table(doc, rows):
    """rows: list of (label, value) tuples"""
    table = doc.add_table(rows=len(rows)+1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ["Metric", "Value"]):
        cell.text = txt
        set_cell_bg(cell, "1E3A5F")
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)
    # Data rows
    for i, (label, value) in enumerate(rows):
        row = table.rows[i+1].cells
        row[0].text = label
        row[1].text = str(value)
        bg = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        for cell in row:
            set_cell_bg(cell, bg)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()

# ── Build Document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default font
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ── TITLE PAGE ──────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Tendertrace Dataset\nExploratory Data Analysis")
tr.bold = True
tr.font.size = Pt(26)
tr.font.color.rgb = RGBColor(30, 58, 95)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Australian Government Procurement — Data Quality & Distribution Report")
sr.font.size = Pt(13)
sr.font.color.rgb = RGBColor(80, 80, 80)
sr.italic = True

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = date_p.add_run(datetime.date.today().strftime("%B %Y"))
dr.font.size = Pt(11)
dr.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ── 1. EXECUTIVE SUMMARY ─────────────────────────────────────────────────────
heading(doc, "1. Executive Summary", level=1)
body(doc,
    "This report presents an exploratory data analysis (EDA) of the Tendertrace export dataset, "
    "comprising Australian government procurement contracts published across federal, state, and "
    "territory jurisdictions. The analysis covers dataset completeness, contract value distributions, "
    "temporal trends, publisher profiles, procurement method breakdowns, and feature correlations — "
    "providing the empirical grounding for the ML-based contract price prediction system."
)
doc.add_paragraph()
add_stat_table(doc, [
    ("Total records (raw export)",     "1,048,575"),
    ("Valid contracts (value > 0)",    "1,009,144"),
    ("Total columns",                  "36"),
    ("Date range",                     "2010 – 2025"),
    ("Median contract value",          "AUD $73,600"),
    ("Mean contract value",            "AUD $4,877,212"),
    ("Max contract value",             "AUD $19.3 billion"),
    ("Distinct publishers (all)",      "1,322 agencies"),
    ("Publishers used in modelling",   "514 (≥50 contracts)"),
    ("UNSPSC category codes",          "534 unique codes"),
    ("Parent category codes",          "55 unique codes"),
    ("Training median duration",       "312 days (~10 months)"),
])

# ── 2. DATA COMPLETENESS ─────────────────────────────────────────────────────
heading(doc, "2. Data Completeness", level=1)
body(doc,
    "The Tendertrace export contains 36 columns sourced from AusTender and various state procurement "
    "portals. Field completeness varies significantly: several administrative columns are largely "
    "unpopulated, while the pre-award fields used for modelling are mostly present. The chart below "
    "shows the percentage of missing values per column."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/01_missing_values.png",
           "Figure 1. Missing values by column (red = >30%, orange = 10–30%, blue = <10%)",
           width=6.2)

heading(doc, "Key completeness observations:", level=3)
bullet(doc, "publisher_company_number, closing_date, amendment_date are >99% missing — these are not used in modelling.", "Critical gaps: ")
bullet(doc, "supplier_state, supplier_name, is_consultancy_services are 30–55% missing — the consultancy flag in particular is unreliable, reducing its usefulness as a feature.", "Partially populated: ")
bullet(doc, "Core modelling fields (publisher_name, publisher_gov_type, category_code, procurement_method, value, contract_start, contract_end) are largely complete (<5% missing).", "Well-populated: ")
bullet(doc, "duration_days is not stored in the raw export — it is derived at ingestion by computing the difference between contract_start and contract_end, with missing dates imputed to the training median of 312 days.", "Derived feature: ")

# ── 3. CONTRACT VALUE DISTRIBUTION ──────────────────────────────────────────
heading(doc, "3. Contract Value Distribution", level=1)
body(doc,
    "Contract values span more than ten orders of magnitude — from under $1 to $19.3 billion — "
    "exhibiting strong right skew. The log-transformed distribution is approximately unimodal, "
    "peaking around $10,000–$100,000 (10⁴–10⁵ AUD). The model therefore predicts log₁p(value) and "
    "transforms back at inference, which stabilises the regression target and prevents extreme outliers "
    "from dominating the loss function."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/02_value_distribution.png",
           "Figure 2. Left: log₁₀ distribution of contract values with median annotated. Right: size bucket breakdown.",
           width=6.2)

heading(doc, "Size bucket breakdown:", level=3)
bullet(doc, "43.0% of contracts are Small (<$50K) — mostly routine operational purchases.", "Small (<$50K): ")
bullet(doc, "37.6% — the model's strongest performance tier.", "Medium ($50K–$500K): ")
bullet(doc, "15.5% — infrastructure, IT, and professional services.", "Large ($500K–$5M): ")
bullet(doc, "3.9% by count but representing the vast majority of total government spend.", "Very Large (>$5M): ")

body(doc,
    "This class imbalance means that a naive model predicting the global median ($73,600) would "
    "appear reasonable on count-weighted metrics but would fail badly on the high-value contracts "
    "that carry the most fiscal risk. The bucket classification layer in the pipeline addresses "
    "this explicitly."
)

# ── 4. TEMPORAL TRENDS ──────────────────────────────────────────────────────
heading(doc, "4. Temporal Trends", level=1)
body(doc,
    "The dataset spans 2010–2025, with contract volume growing substantially over the period. "
    "Annual contract count and median value are shown together below to reveal the interplay "
    "between procurement frequency and average contract size."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/03_contracts_over_time.png",
           "Figure 3. Annual contract count (bars, left axis) and median contract value in AUD K (line, right axis).",
           width=6.2)

heading(doc, "Notable patterns:", level=3)
bullet(doc, "Contract volume grew from ~28K/year in 2010 to a peak of ~112K in 2021–22, before declining to ~63K in 2025 (partial year in export).", "Volume growth: ")
bullet(doc, "Median contract value peaked at ~$110K around 2018 then compressed back to ~$60–70K, likely reflecting an increase in smaller routine contracts being reported in later years.", "Value compression: ")
bullet(doc, "The 2019–2020 spike in median value aligns with COVID-19 emergency procurement (large medical supply and ICT contracts).", "2018–2020 spike: ")
bullet(doc, "The training set was split 80/20 by random sampling, not by time. The time breakdown (2020–22, 22–24, 24–26) in the ML evaluation therefore captures model performance across different market conditions.", "Training split note: ")

# ── 5. TOP PUBLISHERS ────────────────────────────────────────────────────────
heading(doc, "5. Publisher Coverage", level=1)
body(doc,
    "The dataset includes 1,322 distinct publishing agencies. The top 20 publishers by contract count "
    "and total value are shown below. Department of Defence dominates both dimensions substantially."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/04_top_publishers.png",
           "Figure 4. Top 20 publishers by contract count (left) and total contract value (right).",
           width=6.4)

body(doc,
    "The Department of Defence accounts for approximately 180,000+ contracts and over $3 trillion in "
    "cumulative contract value — reflecting both the volume of routine procurement and the presence "
    "of major defence capability contracts worth billions each. Departments of Education and Health are "
    "high in count but modest in individual contract value (primarily grants and service agreements). "
    "The publisher_name feature is one of the strongest predictors in the model: knowing which agency "
    "is procuring provides substantial information about likely contract scale."
)

# ── 6. JURISDICTION ──────────────────────────────────────────────────────────
heading(doc, "6. Jurisdiction Analysis", level=1)
body(doc,
    "The dataset is dominated by Federal contracts (782K records, 77% of total), reflecting AusTender's "
    "comprehensive federal reporting requirements. State coverage is uneven — Queensland has the largest "
    "state presence (145K), while ACT, SA, NT, and TAS each have under 6K records."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/05_value_by_jurisdiction.png",
           "Figure 5. Contract count by jurisdiction (left) and median contract value (right).",
           width=6.2)

body(doc,
    "A key finding is that median contract value is not proportional to volume. South Australia has the "
    "highest median (~$780K), followed by NSW (~$340K) and VIC (~$275K). The Federal median is "
    "relatively low (~$80K) because Defence's high volume of small routine contracts dilutes the median. "
    "Queensland's very low median (~$20K) reflects many small grant and community service contracts. "
    "These jurisdictional differences mean that publisher_gov_type is a useful signal for the model "
    "when combined with other features."
)

# ── 7. PROCUREMENT METHOD ────────────────────────────────────────────────────
heading(doc, "7. Procurement Method", level=1)
body(doc,
    "The procurement_method field contains 62 distinct raw values across jurisdictions, reflecting "
    "inconsistent naming conventions (e.g. 'open', 'open competitive', 'open tender', 'open advertisement' "
    "all describe the same process). For analysis these have been simplified into 8 grouped categories."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/06_procurement_method.png",
           "Figure 6. Contract count (left) and median value (right) by simplified procurement method.",
           width=6.2)

bullet(doc, "Limited/Closed and Open are the most common methods (~350K and ~330K contracts respectively).", "Volume leaders: ")
bullet(doc, "Contracts recorded simply as 'Tender' are rare but have the highest median value (~$600K), capturing large formal competitive procurements.", "Highest value: ")
bullet(doc, "Direct/Sole Source contracts have a median of ~$150K, suggesting agencies use this route for mid-sized specialist engagements.", "Direct sourcing: ")
bullet(doc, "~23% of records fall into Other/Unknown — a data quality gap that limits the feature's predictive power for those records.", "Unknown gap: ")

# ── 8. CONTRACT DURATION ─────────────────────────────────────────────────────
heading(doc, "8. Contract Duration", level=1)
body(doc,
    "Duration is derived at ingestion from contract_start and contract_end timestamps and capped at "
    "3,650 days (10 years) to exclude data-entry errors. It is one of the two strongest predictors in "
    "the model (correlation with log value = 0.52)."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/07_duration_distribution.png",
           "Figure 7. Duration distribution (left) and contracts by duration bucket (right).",
           width=6.2)

body(doc,
    "The distribution is bimodal with peaks at <3 months (221K contracts — likely purchase orders and "
    "spot buys) and 6–12 months (215K — annual service agreements). The visible spikes at exactly 365, "
    "730, and 1095 days in the raw histogram indicate that many contract dates are entered as round "
    "1/2/3-year periods rather than precise calendar dates. The training median of 312 days is used "
    "to impute duration when contract dates are unavailable at inference."
)

# ── 9. CORRELATION MATRIX ────────────────────────────────────────────────────
heading(doc, "9. Feature Correlations", level=1)
body(doc,
    "To understand linear relationships, a Pearson correlation matrix was computed on a 150,000-row "
    "sample. Categorical features were ordinal-encoded prior to correlation computation. Pearson "
    "correlation captures monotonic linear associations; non-linear relationships (which tree models "
    "capture well) are not reflected here."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/08_correlation_matrix.png",
           "Figure 8. Pearson correlation matrix on a 150K-row sample (categoricals ordinal-encoded).",
           width=5.5)

add_stat_table(doc, [
    ("Log Value ↔ Duration",      "0.52  (strong positive — longer contracts cost more)"),
    ("Log Value ↔ Disposition",   "-0.14  (grants tend to be lower value)"),
    ("Log Value ↔ Year",          "-0.03  (no meaningful time trend in value)"),
    ("Proc Method ↔ Jurisdiction","0.48  (naming conventions vary by jurisdiction)"),
    ("Disposition ↔ Year",        "0.27  (grant reporting increased over time)"),
    ("All others",                "≤ 0.16  (weak linear relationship)"),
])
body(doc,
    "The weak linear correlations for most categorical features (≤0.14 with log value) confirm that "
    "the predictive signal is non-linear and interaction-driven — explaining why gradient-boosted "
    "tree models (XGBoost, LightGBM) substantially outperform linear baselines on this dataset. "
    "Duration is the single most informative individual feature."
)

# ── 10. COFOG ANALYSIS ───────────────────────────────────────────────────────
heading(doc, "10. COFOG Government Level", level=1)
body(doc,
    "The publisher_cofog_level field classifies agencies by their functional government classification "
    "(UN COFOG standard). Level 1 represents general government bodies; Level 2 covers state-owned "
    "enterprises. The field is only partially populated (~30% unknown)."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/09_value_by_cofog.png",
           "Figure 9. Median and mean contract value by COFOG government level.",
           width=5.8)

body(doc,
    "Median values are similar across all COFOG levels (~$66–79K), suggesting the classification "
    "alone is not a strong differentiator. However, Level 1 has a dramatically inflated mean "
    "($10.8M vs $1.8M for Level 2) driven by a small number of very large federal defence and "
    "infrastructure contracts. The RAG system uses COFOG codes to provide the agent with functional "
    "government context when interpreting a procurement description."
)

# ── 11. UNSPSC CATEGORIES ────────────────────────────────────────────────────
heading(doc, "11. UNSPSC Category Distribution", level=1)
body(doc,
    "Category codes follow the UNSPSC (United Nations Standard Products and Services Code) taxonomy. "
    "The dataset contains 534 distinct codes with sufficient frequency for modelling (≥50 records). "
    "The top 15 by contract count are shown below."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/10_top_categories.png",
           "Figure 10. Top 15 UNSPSC category codes by contract count, with median value overlay.",
           width=6.2)

body(doc,
    "A significant quality issue is that 'unknown' is the largest single category — indicating that "
    "many agencies do not populate the category code field. Among known codes, 80111600 (staffing and "
    "temporary services) is the most frequent but has a moderate median (~$60K). Code 81110000 "
    "(management advisory services) carries the highest median (~$400K+), reflecting the premium for "
    "high-end consultancy. The RAG domain store indexes the full UNSPSC codeset so the agent can "
    "resolve ambiguous commodity descriptions to the correct code."
)

# ── 12. CONSULTANCY ──────────────────────────────────────────────────────────
heading(doc, "12. Consultancy Services", level=1)
body(doc,
    "The is_consultancy_services flag is intended to distinguish advisory/consultancy contracts from "
    "goods and services procurement. However, it is poorly populated — only 'Non-Consultancy' and "
    "'Unknown' appear as effective values in the cleaned export, with no records reliably marked "
    "as 'Consultancy'. This is a known data quality limitation of the AusTender feed for state "
    "jurisdictions."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/11_consultancy_split.png",
           "Figure 11. Contract count, median value, and total value for consultancy flag values.",
           width=6.0)

body(doc,
    "Non-Consultancy contracts (600K records, $4 trillion total) have a slightly higher median "
    "($82K vs $50K for Unknown). Despite the field's limitations, it is retained as a model feature "
    "because even the binary non-consultancy/unknown signal carries some predictive value."
)

# ── 13. JURISDICTION × YEAR HEATMAP ─────────────────────────────────────────
heading(doc, "13. Jurisdiction × Year Activity Heatmap", level=1)
body(doc,
    "The heatmap below cross-tabulates contract volume by jurisdiction and year (2015–2024), "
    "revealing how data coverage has evolved over time and across reporting bodies."
)
doc.add_paragraph()
add_figure(doc, f"{IMG}/12_jurisdiction_year_heatmap.png",
           "Figure 12. Contract count by jurisdiction and year (2015–2024).",
           width=6.4)

bullet(doc, "Federal volume has been consistently 39K–71K contracts per year, the backbone of the dataset.", "Federal consistency: ")
bullet(doc, "QLD saw a sharp surge in 2019–2022 (up to 33,896 contracts in 2022) followed by a dramatic drop in 2023–24, likely reflecting a change in state reporting practice or data export scope.", "QLD anomaly: ")
bullet(doc, "NT data only begins from 2020 onward, suggesting this jurisdiction was added to the Tendertrace feed later.", "NT late arrival: ")
bullet(doc, "ACT, TAS, SA, VIC, WA have shown steady modest growth over the period.", "State growth: ")
body(doc,
    "The temporal unevenness across jurisdictions is important context for the ML evaluation: "
    "the model is trained and tested on a random 80/20 split across all years and jurisdictions, "
    "so its performance reflects average behaviour rather than out-of-time generalisability."
)

# ── 14. IMPLICATIONS FOR MODELLING ──────────────────────────────────────────
heading(doc, "14. Implications for the ML Pipeline", level=1)
body(doc,
    "The EDA findings directly shaped the modelling and system design decisions:"
)
bullet(doc, "Contract value spans 10+ orders of magnitude with strong right skew — modelled as log₁p(value) to stabilise variance.", "Log-transform target: ")
bullet(doc, "Duration (r=0.52 with log value) and publisher_name are the strongest individual predictors.", "Key features: ")
bullet(doc, "Categoricals show weak linear correlations but strong non-linear interactions captured by XGBoost/LightGBM.", "Tree models: ")
bullet(doc, "Rare category values (appearing <50 times) are collapsed to '<col>_other' to prevent overfitting on low-frequency labels.", "Rare-category collapse: ")
bullet(doc, "publisher_portfolio and publisher_cofog_level are auto-filled at inference from a lookup table built during training, so users only need to provide publisher_name.", "Inference simplification: ")
bullet(doc, "The RAG domain store covers UNSPSC commodity codes, COFOG functional classification, and AusTender valid values — addressing the three main fields where free-text input must be mapped to structured codes.", "RAG coverage: ")
bullet(doc, "The 'unknown' dominance in category_code and consultancy fields means the model must perform well even when these fields are missing — which native categorical encoding with an _other bucket handles gracefully.", "Graceful degradation: ")

# ── SAVE ─────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"Saved: {OUT}")
